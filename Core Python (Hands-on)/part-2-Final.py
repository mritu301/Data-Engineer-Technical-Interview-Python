# -*- coding: utf-8 -*-
"""
Created on Sat May 30 10:08:08 2020

@author: MJ
"""

import docx
import os

#To fixed the root file, so that navigation will be smooth
ROOT_PATH = os.path.dirname(os.path.abspath('__file__'))

#----- Below are the lists of methods wh
# Creating a dictionary for using 2 lists
def createDictionary(lst1, lst2):
    Dict = {}
    counter = 0 
    for value in lst1:
        Dict[lst2[counter] + ":"] = value
        counter = counter + 1
        #print(Dict)
    return Dict

def readFile(filename):
    lst = []
    with open(filename, 'r') as file:
        lines = file.readlines()
        for line in lines:
            temp = (line.strip()).split("|")
            lst.append(temp)
            #print(lst)
    file.close()
    return lst

def flatten(A):
    rt = []
    for i in A:
        if isinstance(i,list): rt.extend(flatten(i))
        else: rt.append(i)
    return rt

fileName = ROOT_PATH + '/key_attributes.txt'
print(flatten(readFile(fileName)))


finalTemplate = ROOT_PATH + "\\final_template" + ".docx"

#print(finalTemplate)

#Read the key_attributes text file and store it into keyAttributes list
with open(ROOT_PATH + '/key_attributes.txt', 'r') as file1:
    lines = file1.readlines()
    for line in lines:
        keyAttributes = line.split("|")
    print(keyAttributes)
file1.close()

#Create a trade attribuites keys for mapping
tradeAttributes = ["FCM TERM SHEET",
                   "Bank Ref",  
                   "Trade Date", 
                   "Initial Evaluation Date", 
                   "Effective Date", 
                   "Notional Amount", 
                   "Fixed Rate"]

# Creating a key attribuites dictionary for accessing the value on the basis of trade attributes
keyAttributesDict = createDictionary(keyAttributes, tradeAttributes)
#print(keyAttributesDict)

#Read the payment_date text file and store it into paymentDate list
paymentDate = []
with open(ROOT_PATH + '/payment_date.txt', 'r') as file2:
    lines = file2.readlines()
    for line in lines:
        lst = (line.strip()).split("|")
        paymentDate.append(lst)
    print(paymentDate)
file2.close()

#paymentDateDict = createDictionary(keyAttributes, tradeAttributes)
#print(keyAttributesDict)

doc = docx.Document(finalTemplate)
print(doc)


print(len(doc.tables))                                                         # To see the numbers of tables in a doc

table1 = doc.tables[0]
table2 = doc.tables[1]

#print(len(table1.rows))                                                       # To see the length of rows in table1
#print(len(table1.columns))                                                    # To see the length of columns in table1
#print(len(table2.rows))                                                       # To see the length of rows in table2
#print(len(table2.columns))                                                    # To see the length of columns in table2

#editing the table 1 using key attribuies (dictionary)
for r in range(len(table1.rows)):
    for c in range(len(table1.columns)): 
        #print(doc.tables[0].cell(r,c).text)
        keysList = keyAttributesDict.keys()
        k = doc.tables[0].cell(r,c).text
        #print("k = ", k)
        #print("keysList = ", keysList)
        if (k in keysList):
            #print("k -", k, " and value - ", keyAttributesDict[k])
            doc.tables[0].cell(r,c + 1).text = keyAttributesDict[k]
            
'''           
for r in range(len(table2.rows)):
    for c in range(len(table2.columns)): 
        #print(doc.tables[1].cell(r,c).text)
        k = doc.tables[1].cell(r,c).text
        #print("k = ", k)
'''


#editing the table 2 using payment date list
for rec in range(0, len(paymentDate), 1):
    if rec == 0:
        doc.tables[1].cell(1,0).text = paymentDate[rec][0]
        doc.tables[1].cell(1,1).text = paymentDate[rec][1]
    else:
        cells  = doc.tables[1].add_row().cells
        cells[0].text = paymentDate[rec][0]
        cells[1].text = paymentDate[rec][1]
    
#doc.tables[1].rows(1).Delete()

'''
tempTableName = ""
counter = 1
for table in doc.tables:
    for row in table.rows:
        print(row.cells)
        for cell in row.cells:
            #print("table - ", table, "table cells details - ", table.cell, "cell - ", cell, " and text -", cell.text)
            if(tempTableName != table):
                print("table ", counter,  " - ", table)
                counter = counter + 1
            
            tempTableName = table
            keysList = keyAttributesDict.keys()
            
            #if cell.text in keysList:
                
            print("text -", cell.text)
  
'''          
     
finalTemplate1 = ROOT_PATH + "\\final_template1" + ".docx"
doc.save(finalTemplate1)
os.system("start final_template1.docx")
