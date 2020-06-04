# -*- coding: utf-8 -*-
"""
Created on Sat May 30 13:26:49 2020

@author: MJ
"""

import docx
import os

#To fixed the root file, so that navigation will be smooth
ROOT_PATH = os.path.dirname(os.path.abspath('__file__'))

#----- Below are the lists of methods which is being used for 2nd part (File and Doc merging)

#Read the word document using docx library
def readWordDocument(documentName):
    return docx.Document(documentName)

# Creating a dictionary for using 2 lists
def createDictionary(lst1, lst2):
    Dict = {}
    counter = 0 
    for value in lst1:
        Dict[lst2[counter] + ":"] = value
        counter = counter + 1
        #print(Dict)
    return Dict

#To read the key_attributes and payment_date text files and split the data on the 
#basis of '|' and return it in to list
def readFile(filename):
    lst = []
    with open(filename, 'r') as file:
        lines = file.readlines()
        for line in lines:
            temp = (line.strip()).split("|")
            lst.append(temp)
            #print(lst)
    file.close()
    return lst


#This below method is to flatten the string arrays into single array
def flatten(A):
    rt = []
    for i in A:
        if isinstance(i,list): rt.extend(flatten(i))
        else: rt.append(i)
    return rt

#editing the table 1 using key attribuies (dictionary)
def table1Modification(doc, keyAttributesDict):
    for r in range(len(doc.tables[0].rows)):
        for c in range(len(doc.tables[0].columns)): 
            keysList = keyAttributesDict.keys()
            k = doc.tables[0].cell(r,c).text
            if (k in keysList):
                doc.tables[0].cell(r,c + 1).text = keyAttributesDict[k]
    return doc

#editing the table 2 using payment date list
def table2Modification(doc, paymentDate):
    for rec in range(0, len(paymentDate), 1):
        if rec == 0:
            doc.tables[1].cell(1,0).text = paymentDate[rec][0]
            doc.tables[1].cell(1,1).text = paymentDate[rec][1]
        else:
            cells  = doc.tables[1].add_row().cells
            cells[0].text = paymentDate[rec][0]
            cells[1].text = paymentDate[rec][1]
    return doc
        
#Save the modified document in specified path (document name should be fully qualified name)
def saveDocument(doc, documentName):
    doc.save(documentName)
    
    
    
#---- Using the above methods, to complete the 2nd part of files and document merging

#Read the key_attributes text file
attributesFileName = ROOT_PATH + '/key_attributes.txt'
keyAttributes = flatten(readFile(attributesFileName))

#Read the payment_date text file
paymentDateFileName = ROOT_PATH + '/payment_date.txt'
paymentDate = readFile(paymentDateFileName)

#print(keyAttributes)
#print(paymentDate)

#Create a trade attribuites keys for mapping
tradeAttributes = ["FCM TERM SHEET",
                   "Bank Ref",  
                   "Trade Date", 
                   "Initial Evaluation Date", 
                   "Effective Date", 
                   "Notional Amount", 
                   "Fixed Rate"]

# Creating a key attribuites dictionary for accessing the value on the basis of trade attributes
keyAttributesDict = createDictionary(keyAttributes, tradeAttributes)
#print(keyAttributesDict)

#Read the final_template docx
finalTemplate = ROOT_PATH + "\\final_template" + ".docx"
print(finalTemplate)
doc = readWordDocument(finalTemplate)
print(doc)


print(len(doc.tables))                                                         # To see the numbers of tables in a doc
doc = table1Modification(doc, keyAttributesDict)                               # To merge the key attributies into table 1
doc = table2Modification(doc, paymentDate)                                     # To merge the list of payment date into table 2
doc = saveDocument(doc, finalTemplate)                                         # Save the modified document
os.system("start final_template.docx")                                         # Open the modified document with the help of os